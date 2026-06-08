import re
from dataclasses import dataclass
from pathlib import Path

import fitz
import pandas as pd
from docx import Document


@dataclass
class ParsedPage:
    text: str
    page_number: int | None
    section_title: str | None = None


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return False
    return bool(re.match(r"^(\d+(\.\d+)*\.?\s+)?[A-Z][A-Za-z0-9 ,:&()/-]{2,}$", stripped))


def _latest_heading(text: str) -> str | None:
    for line in reversed(text.splitlines()):
        if _looks_like_heading(line):
            return line.strip()
    return None


def parse_pdf(path: Path) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    try:
        with fitz.open(path) as doc:
            for index, page in enumerate(doc, start=1):
                text = clean_text(page.get_text("text"))
                if text:
                    pages.append(ParsedPage(text=text, page_number=index, section_title=_latest_heading(text)))
    except Exception as exc:
        raise ValueError("The PDF could not be read. It may be corrupted, encrypted, or scanned without text.") from exc
    return pages


def parse_docx(path: Path) -> list[ParsedPage]:
    doc = Document(path)
    blocks: list[str] = []
    current_heading: str | None = None
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        style = ((paragraph.style.name if paragraph.style else "") or "").lower()
        if "heading" in style or _looks_like_heading(text):
            current_heading = text
        blocks.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"Table {table_index}\n" + "\n".join(rows))

    content = "\n\n".join(blocks)
    return [ParsedPage(text=content, page_number=None, section_title=current_heading)] if content else []


def parse_txt(path: Path) -> list[ParsedPage]:
    text = clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    return [ParsedPage(text=text, page_number=None, section_title=_latest_heading(text))] if text else []


def parse_csv(path: Path) -> list[ParsedPage]:
    frame = pd.read_csv(path)
    rows = []
    for index, row in frame.iterrows():
        values = [f"{column}: {row[column]}" for column in frame.columns if pd.notna(row[column])]
        rows.append(f"Row {index + 1}. " + "; ".join(values))
    text = clean_text("\n".join(rows))
    return [ParsedPage(text=text, page_number=None, section_title="CSV Data")] if text else []


def parse_document(path: Path, filename: str) -> list[ParsedPage]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".txt":
        return parse_txt(path)
    if suffix == ".csv":
        return parse_csv(path)
    raise ValueError(f"Unsupported file type for {filename}")
